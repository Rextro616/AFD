import pandas as pd
from docx import Document
from bs4 import BeautifulSoup

def readFile(pathFile, extension):

    if extension == 'csv':
        return readCSV(pathFile)
    elif extension == 'xlsx':
        return readExcel(pathFile)
    elif extension == 'docx':
        return readDocument(pathFile)
    elif extension == 'html':
        return readHtml(pathFile)
    else:
        return "Tipo de archivo no soportado."

def readCSV(pathFile):
    try:
        df = pd.read_csv(pathFile)
        loops = [] 
        loops = [(f"{row_index}, {col_index}", str(value))
            for row_index, row in df.iterrows()
            for col_index, value in row.items() 
            if isinstance(value, str) and value]  # Filtrar solo valores de tipo string que no estén vacíos
        return loops  
    except Exception as e:
        raise Exception(f"No se pudo leer el archivo CSV: {e}")

def readDocument(pathFile):
    try:
        doc = Document(pathFile)
        fullText = [paragraphs.text for paragraphs in doc.paragraphs]
        return '\n'.join(fullText)
    except Exception as e:
        return f"Error al leer el archivo DOCX: {e}"

def readExcel(pathFile):
    try:
        df = pd.read_excel(pathFile)
        loops = [] 
        loops = [(f"{row_index}, {col_index}", str(value))
            for row_index, row in df.iterrows()
            for col_index, value in row.items()
            if isinstance(value, str) and value]
        return loops  
    except Exception as e:
        raise Exception(f"No se pudo leer el archivo Excel: {e}")
    

def readHtml(pathFile):
    try:
        with open(pathFile, 'r', encoding='utf-8') as file:
            soup = BeautifulSoup(file, 'html.parser')
        try:
            tables = pd.read_html(pathFile)
            return tables[0].to_string()  # Mostramos solo la primera tabla
        except:
            pass        
        return soup.get_text()
    except Exception as e:
        return f"Error al leer el archivo HTML: {e}"
